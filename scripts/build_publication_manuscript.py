"""Create the journal-neutral SeismoFlux manuscript as an editable DOCX."""

from __future__ import annotations

import argparse
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


INK = "172235"
MUTED = "5D6B7E"
LINE = "DCE4EC"
PALE = "F3F6F8"
PALE_BLUE = "EAF2F6"
PALE_ORANGE = "FFF1EB"
BLUE = "1F6F8B"
ORANGE = "D85A30"
GREEN = "168F73"


def _set_east_asia(run, typeface: str) -> None:
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), typeface)


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(cell, *, top: int = 90, start: int = 110, bottom: int = 90, end: int = 110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def _set_cell_text(cell, text: str, *, bold: bool = False, color: str = INK, size: float = 9.0) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    _set_east_asia(run, "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def _add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    run.font.name = "Arial"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    tail = paragraph.add_run(" 页")
    tail.font.name = "Microsoft YaHei"
    tail.font.size = Pt(8.5)
    tail.font.color.rgb = RGBColor.from_string(MUTED)


def _set_run_font(run, *, latin: str = "Times New Roman", east_asia: str = "宋体", size: float = 10.5, color: str = INK, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = latin
    _set_east_asia(run, east_asia)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _add_paragraph(
    document: Document,
    text: str = "",
    *,
    style: str | None = None,
    first_line: bool = True,
    align: WD_ALIGN_PARAGRAPH | None = WD_ALIGN_PARAGRAPH.JUSTIFY,
    space_after: float = 5.5,
) :
    paragraph = document.add_paragraph(style=style)
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_after = Pt(space_after)
    if first_line:
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
    if text:
        run = paragraph.add_run(text)
        _set_run_font(run)
    return paragraph


def _add_rich_paragraph(document: Document, parts: Sequence[tuple[str, bool, bool]], *, first_line: bool = True):
    paragraph = _add_paragraph(document, first_line=first_line)
    for text, bold, italic in parts:
        run = paragraph.add_run(text)
        _set_run_font(run, bold=bold, italic=italic)
    return paragraph


def _add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    _set_run_font(
        run,
        latin="Arial",
        east_asia="Microsoft YaHei",
        size=14.5 if level == 1 else 12.0,
        color=INK,
        bold=True,
    )


def _add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Cm(0.72)
    paragraph.paragraph_format.first_line_indent = Cm(-0.36)
    paragraph.paragraph_format.space_after = Pt(3.5)
    run = paragraph.add_run(text)
    _set_run_font(run)


def _add_key_box(document: Document, title: str, body: str, *, accent: str = ORANGE, fill: str = PALE_ORANGE) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(0.42)
    table.columns[1].width = Cm(16.2)
    _set_cell_shading(table.cell(0, 0), accent)
    _set_cell_shading(table.cell(0, 1), fill)
    _set_cell_margins(table.cell(0, 0), top=80, start=40, bottom=80, end=40)
    _set_cell_margins(table.cell(0, 1), top=130, start=170, bottom=130, end=170)
    _set_cell_text(table.cell(0, 0), "")
    cell = table.cell(0, 1)
    cell.text = ""
    title_p = cell.paragraphs[0]
    title_p.paragraph_format.space_after = Pt(3)
    title_run = title_p.add_run(title)
    _set_run_font(title_run, latin="Arial", east_asia="Microsoft YaHei", size=10.5, color=accent, bold=True)
    body_p = cell.add_paragraph()
    body_p.paragraph_format.space_after = Pt(0)
    body_p.paragraph_format.line_spacing = 1.25
    body_run = body_p.add_run(body)
    _set_run_font(body_run, size=10.0)
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def _add_figure(document: Document, path: Path, caption: str, *, width_cm: float = 16.8) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    caption_p = document.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    caption_p.paragraph_format.keep_with_next = True
    caption_p.paragraph_format.space_after = Pt(7)
    caption_p.paragraph_format.line_spacing = 1.1
    label, _, rest = caption.partition(" ")
    label_run = caption_p.add_run(label + " ")
    _set_run_font(label_run, latin="Arial", east_asia="Microsoft YaHei", size=9.0, bold=True)
    rest_run = caption_p.add_run(rest)
    _set_run_font(rest_run, size=9.0, color=MUTED)


def _add_table(
    document: Document,
    headers: Sequence[str],
    rows: Iterable[Sequence[str]],
    *,
    widths_cm: Sequence[float] | None = None,
    font_size: float = 8.7,
) -> None:
    rows = list(rows)
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = widths_cm is None
    header = table.rows[0]
    _set_repeat_table_header(header)
    for index, value in enumerate(headers):
        cell = header.cells[index]
        _set_cell_shading(cell, BLUE)
        _set_cell_margins(cell)
        _set_cell_text(cell, value, bold=True, color="FFFFFF", size=font_size)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths_cm is not None:
            cell.width = Cm(widths_cm[index])
    for row_index, values in enumerate(rows):
        row = table.add_row()
        _prevent_row_split(row)
        for index, value in enumerate(values):
            cell = row.cells[index]
            _set_cell_margins(cell)
            if row_index % 2:
                _set_cell_shading(cell, "F8FAFC")
            _set_cell_text(cell, str(value), size=font_size)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths_cm is not None:
                cell.width = Cm(widths_cm[index])
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)


def _configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.line_spacing = 1.5

    for level, size in ((1, 14.5), (2, 12.0), (3, 11.0)):
        style = document.styles[f"Heading {level}"]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(INK)
        style.paragraph_format.space_before = Pt(11 if level == 1 else 7)
        style.paragraph_format.space_after = Pt(5)

    if "Figure Caption" not in [style.name for style in document.styles]:
        style = document.styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(9)
        style.font.color.rgb = RGBColor.from_string(MUTED)


def _configure_sections(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.9)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_p.paragraph_format.space_after = Pt(2)
    run = header_p.add_run("SeismoFlux  ·  原始研究论文稿")
    _set_run_font(run, latin="Arial", east_asia="Microsoft YaHei", size=8.5, color=MUTED, bold=True)
    p_pr = header_p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), LINE)
    borders.append(bottom)
    p_pr.append(borders)

    _add_page_number(section.footer.paragraphs[0])


def _title_page(document: Document) -> None:
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(16)
    type_p = document.add_paragraph()
    type_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = type_p.add_run("ORIGINAL ARTICLE  ·  RETROSPECTIVE DEVELOPMENT EVIDENCE")
    _set_run_font(run, latin="Arial", east_asia="Microsoft YaHei", size=9.5, color=BLUE, bold=True)
    type_p.paragraph_format.space_after = Pt(13)

    title_p = document.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(8)
    title_run = title_p.add_run("近期地震活动相对长期核密度背景的回顾性区域排序增益")
    _set_run_font(title_run, latin="Arial", east_asia="Microsoft YaHei", size=18.5, color=INK, bold=True)

    subtitle_p = document.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(16)
    subtitle_run = subtitle_p.add_run("——中国大陆M5–6地震的严格时间前推滚动历史检验")
    _set_run_font(subtitle_run, latin="Arial", east_asia="Microsoft YaHei", size=13.5, color=MUTED, bold=False)

    english_p = document.add_paragraph()
    english_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    english_p.paragraph_format.space_after = Pt(17)
    english_run = english_p.add_run(
        "Retrospective Regional-Ranking Gain from Recent Seismicity over a Long-Term "
        "Smoothed-Seismicity Background for M5–6 Earthquakes in Mainland China"
    )
    _set_run_font(english_run, latin="Arial", east_asia="Microsoft YaHei", size=11.5, color=BLUE, bold=True)

    author_p = document.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_run = author_p.add_run("作者与单位：待作者确认")
    _set_run_font(author_run, latin="Arial", east_asia="Microsoft YaHei", size=10.0, color=MUTED)
    author_p.paragraph_format.space_after = Pt(20)

    _add_key_box(
        document,
        "核心结果",
        "在30天、600,000 km²报警面积上限下，近期30天地震活动与长期背景的冻结混合模型将历史时间前推命中从5/21提高到9/21；新增4个规则化震群，召回提高19.05个百分点。该结果尚不是成熟的真实前瞻证明。",
    )

    label_p = document.add_paragraph()
    label_p.paragraph_format.space_before = Pt(10)
    label_p.paragraph_format.space_after = Pt(3)
    label_run = label_p.add_run("研究状态")
    _set_run_font(label_run, latin="Arial", east_asia="Microsoft YaHei", size=9.2, color=MUTED, bold=True)
    status_p = document.add_paragraph()
    status_p.paragraph_format.space_after = Pt(10)
    status_run = status_p.add_run("严格时间前推的历史开发检验；真实前瞻协议已预登记并授权，但截至2026年8月31日首个合法起报时刻尚未到达，真实起报为0期。")
    _set_run_font(status_run, size=9.5, color=INK)

    version_p = document.add_paragraph()
    version_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version_p.paragraph_format.space_before = Pt(28)
    version_run = version_p.add_run("历史开发回放证据  ·  2026年8月31日")
    _set_run_font(version_run, latin="Arial", east_asia="Microsoft YaHei", size=9.0, color=MUTED)
    document.add_page_break()


def _abstract(document: Document) -> None:
    _add_heading(document, "摘要", 1)
    _add_rich_paragraph(
        document,
        [
            ("目的：", True, False),
            ("检验近期地震活动能否在不提高600,000 km²报警面积上限的条件下，提高未来30天中国大陆M5–6规则化震群的区域召回。", False, False),
            (" 方法：", True, False),
            ("主要比较基于40,898条冻结地震目录、15,697个约25 km等面积网格和三个严格按时间向未来外推的评价折，比较长期地震背景B0与加入近30天地震活动的B0_R30；205期异常报告另作次级增量检验。两模型均输出相对条件强度和网格顺位，以300,000–960,000 km²五档报警面积上限评价；一级终点为30天、600,000 km²、21个按75 km/30天规则合并的M5–6震群。", False, False),
            (" 结果：", True, False),
            ("B0_R30将命中从5/21提高到9/21，召回由23.81%提高到42.86%，增加4群和19.05个百分点。两模型在该面积上限下的平均实际面积分别为599,447和599,666 km²，相差219 km²，小于一个完整网格。2,000次配对震群Bootstrap的95%区间为+4.76至+38.10个百分点，正增益复本比例为0.9905；该比例不是模型正确概率，也不构成传统显著性证明。三折命中为3/8→5/8、2/6→4/6和0/7→0/7。B0_R30在300,000 km²已命中7/21，高于B0在600,000 km²的5/21。四个新增命中均位于西北地区，其中两个落在同一25 km网格。异常快照和动态异常没有增加30天主终点命中。", False, False),
            (" 结论：", True, False),
            ("近期地震活动为该历史样本提供了有实际幅度的区域召回增益，主要机制与地震聚集或活动序列延续一致。由于样本量小、地域集中、第三折无改善，且B0_R30增益尚未接受专门置乱及成熟真实前瞻检验，该结果应作为值得继续前瞻检验的强历史线索，而不是已证实的实际地震预测能力。", False, False),
        ],
        first_line=False,
    )
    keywords = _add_paragraph(document, first_line=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    run = keywords.add_run("关键词：")
    _set_run_font(run, bold=True)
    run = keywords.add_run("地震预测；区域召回；近期地震活动；报警面积上限；时间外推；规则化震群")
    _set_run_font(run)

    _add_heading(document, "Abstract", 1)
    abstract_en = (
        "We tested whether recent seismicity improves regional M5–6 earthquake forecasting in mainland China without raising the alarm-area cap. "
        "A frozen catalogue of 40,898 earthquakes was evaluated on 15,697 approximately 25-km cells using three forward-in-time retrospective folds. "
        "The long-term smoothed-seismicity background (B0) was compared with a fixed mixture that assigned 25% weight to seismicity observed during the preceding 30 days (B0_R30). "
        "At the primary 30-day, 600,000-km² alarm-area cap, recall increased from 5/21 (23.81%) to 9/21 (42.86%) rule-based clusters, a gain of four clusters or 19.05 percentage points; mean realized areas differed by only 219 km². "
        "A 2,000-replicate paired cluster bootstrap yielded a 95% interval of +4.76 to +38.10 percentage points and a positive-gain proportion of 0.9905; this conditional proportion is neither a posterior probability nor a conventional significance test. "
        "The fold-wise hit counts were 3/8 to 5/8, 2/6 to 4/6, and 0/7 to 0/7. "
        "All four additional hits occurred in northwestern China, and two occupied the same 25-km cell. "
        "The evidence therefore supports recent seismicity as a useful retrospective regional-ranking component, while geographic concentration, small sample size, and zero prospective issues as of the manuscript freeze preclude claims of demonstrated operational predictability."
    )
    paragraph = _add_paragraph(document, abstract_en, first_line=False)
    for run in paragraph.runs:
        _set_run_font(run, latin="Times New Roman", east_asia="宋体", size=10.0)
    keywords_en = _add_paragraph(document, first_line=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    run = keywords_en.add_run("Keywords: ")
    _set_run_font(run, bold=True)
    run = keywords_en.add_run("earthquake forecasting; regional recall; recent seismicity; fixed alarm area; temporal extrapolation")
    _set_run_font(run)
    document.add_page_break()


def _introduction(document: Document) -> None:
    _add_heading(document, "1 引言", 1)
    _add_paragraph(
        document,
        "地震发生具有明显的时空聚集，但把这种统计规律转化为稳定、可检验的未来区域预测仍然困难。一个容易产生误解的评价方式，是只报告命中了多少次，却不同时约束报警覆盖范围。报警面积越大，命中通常越多；若不同模型使用不同面积，召回提高可能只是扩大报警范围的结果。Molchan图和报警面积检验强调漏报率与报警时空占比必须共同解释[1–3]。CSEP一类前瞻试验进一步要求模型、输入、评价区间和评分规则在未来地震发生前固定[4]。",
    )
    _add_paragraph(
        document,
        "统计地震学通常把历史地震目录分解为长期空间背景与由既有地震触发的短期活动。ETAS及其空间扩展用条件强度描述这种自激过程[5–7]，近年来的神经点过程和目录编码模型也反复发现，近期地震历史是短期预测的重要信息源[8,9]。然而，复杂模型并不自动等于更好的预测；参考模型有助于识别真实增量[10]，严格时间外推和未来盲检则决定结论能否推广到未知地震[4,12]。",
    )
    _add_paragraph(
        document,
        "SeismoFlux的科学目标不是给出单一震中的确定性预言，而是在预先限定的报警面积上限下，提高未来目标地震所属区域的召回。前期研究同时接入地震目录、异常报告、断层与底图等数据，并尝试ETAS、核密度背景和异常增量模型。当前冻结ETAS实现未达到数值资格，因此没有可比较的ETAS成绩；这不代表ETAS没有科学价值。按照预登记规则，本研究以满足预设数值与时间外推条件的75 km地震核密度背景作为基准模型，先回答一个更窄但可检验的问题：在长期背景上加入近30天地震活动，是否能在相同面积上限下覆盖更多M5–6规则化震群。",
    )
    _add_paragraph(
        document,
        "本文报告冻结历史开发回放的完整结果，并用全部21个规则化震群和四个新增命中案例解释模型行为。所有案例均在总体结果形成后选出，仅用于说明模型行为；统计结论始终来自全部样本。本文还报告异常增量的负结果、主要限制和已经冻结的真实前瞻方案。",
    )


def _data_methods(document: Document, figures: Path) -> None:
    _add_heading(document, "2 数据与方法", 1)
    _add_heading(document, "2.1 数据、用途与未使用信息", 2)
    _add_paragraph(
        document,
        "研究区取中国大陆连续陆域，排除海南、台湾及其他离岛，冻结等面积投影为Albers投影（lat_1=25°、lat_2=47°、lon_0=105°，WGS84），投影面积为9,415,305.754 km²。空间积分和报警选择基于15,697个约25 km的目标无关固定网格；边界网格按研究区裁剪后的实际面积计入。地震目录冻结到2026年7月9日可见记录，异常报告冻结到2026年7月2日可见期次。主要数据及其用途见表1。",
    )
    _add_table(
        document,
        ["数据", "规模", "本研究中的用途", "对最终B0_R30的作用"],
        [
            ("地震目录", "43,785条源记录；去重后40,898事件", "38,375条M3+目录与5,410条M5+目录合并；起报日前可见的M4+用于B0和R30，起报后M5–6仅用于评分", "直接使用"),
            ("异常报告", "205期；59,904条观测", "重建3,217,885行时间可用快照与动态特征，比较三个异常增量模型", "已检验，但主终点无增益"),
            ("固定空间网格", "15,697格", "积分相对强度、排序并形成固定面积报警前缀", "直接使用"),
            ("目标无关构造区", "39区", "异常空间置乱与事后区域稳健性", "不作为预测特征"),
            ("断层与底图", "已接入", "数据盘点与可视化背景；本轮未正式评价其预测增量", "未进入最终模型"),
            ("人工预测地点/时间/震级", "—", "明确禁止作为特征或标签", "未使用"),
        ],
        widths_cm=[2.8, 2.0, 7.4, 4.1],
    )
    _add_paragraph(
        document,
        "目录时间统一到Asia/Shanghai语义。对任一起报时刻T，模型只能读取available_at不晚于T的记录。预测窗定义为(T,T+h]。真实震中不用于生成候选、加密网格或划定区域，只在报警区域冻结后用于命中/漏报评分。原始数据的再分发许可尚未明确，因此公开成果包只包含派生图、汇总表和可复现代码，不附逐行原始目录或异常表。",
    )

    _add_heading(document, "2.2 长期背景与近期活动模型", 2)
    _add_paragraph(
        document,
        "长期背景B0使用T之前可见的1970年以来M4+地震，在冻结等面积投影中进行75 km高斯核平滑。近期活动R30只使用(T−30天,T]内可见的M4+地震。每个成分在研究区内归一化为相对空间质量。三个外折分别只在其历史训练前缀内，用滚动内部验证期M4+目标的平均对数空间密度，从混合权重{0, 0.25, 0.5, 0.75}中选择；容差内并列时取较小权重，三折均选得0.25，因而最终模型为B0_R30=0.75×B0+0.25×R30。模型输出是相对条件强度和网格顺位，不是未来30天的绝对发震概率。",
    )
    _add_paragraph(
        document,
        "作为各数据源增量比较，本研究还比较了B0、B0_R30、报告覆盖控制B0_C、异常快照B0_C_A_snapshot、异常动态B0_C_A_dynamic以及完整组合B0_R30_C_A_dynamic六个简单模型。B0_C在B0偏置上加入当前及过去一年台站/测项报告完整度与报告广度，用于区分“报告更多”和“异常更强”。带特征的模型使用逐起报条件空间多项式似然和ridge正则；预处理、缺失处理、标准化和参数选择仅在每个评价折之前的历史前缀完成。本轮预先限定为可解释的简单模型，以控制开发自由度。",
    )
    _add_figure(
        document,
        figures / "figure_05_method_overview.png",
        "图1 研究设计。起报日前的长期目录形成B0，近30天目录形成R30；冻结混合在15,697个网格上排序，并以相同报警面积上限评价。T之后的目标地震只用于报警冻结后的评分。",
    )

    _add_heading(document, "2.3 时间外推与规则化震群", 2)
    _add_paragraph(
        document,
        "30天一级终点包含24个按固定日历规则选择且互不重叠的起报日。三段评价区间为2023-07至2024-04、2024-07至2025-04和2025-07至2026-04；各段模型仅使用更早历史拟合。为降低同一地震序列被重复计数的影响，研究区M5–6目标按震中距离不超过75 km且时间差不超过30天构图，并以连通分量合并为规则化震群。该规则不证明震群在统计或物理上相互独立。30天窗共有29个M5–6事件，合并为21群，三折分别为8、6和7群；90天次级终点含9个起报日，32个事件合并为22群。",
    )

    _add_heading(document, "2.4 固定面积评价和不确定性", 2)
    _add_paragraph(
        document,
        "每个起报日先按网格质量除以裁剪面积得到相对强度，再按强度降序、row、column和cell ID依次打破并列，随后累加完整裁剪网格面积，形成300,000、450,000、600,000、750,000和960,000 km²五档报警前缀；若下一个网格会超出面积上限即停止，不跳过该格继续挑选。一级终点预先固定为30天、600,000 km²、M5–6规则化震群严格召回。平均实际报警面积使用全部起报日计算，而不是只在有目标的起报日计算。次级指标包括完整面积—召回曲线、90天召回、Molchan曲线和目标点连续空间对数密度。",
    )
    _add_paragraph(
        document,
        "B0_R30与B0的不确定性用根种子147的2,000次配对震群Bootstrap估计，重采样单位为21个规则化震群，保持两个模型对同一震群的成对关系。报告点估计、95%分位区间和正增益复本比例。该比例只描述在已拟合、已选择模型和这21群条件下的重采样结果，不是后验成功概率。该Bootstrap没有重拟合模型，也没有按地域区块重采样，并假定规则化震群可交换，因此不能覆盖参数估计、模型选择和地域集中带来的全部不确定性。",
    )

    _add_heading(document, "2.5 异常增量归因", 2)
    _add_paragraph(
        document,
        "异常快照、全部异常和动态演化在30天主终点的三个预登记增量均为0。进一步使用时间置乱和空间置乱各200个索引，并在三个外折分别重拟合，共得到1,200个折级结果；时间置乱科学拟合失败58/200，空间置乱失败11/200，均超过冻结的5%失败上限。因此异常增量被判为证据不足，而非有效。需要强调的是，这些置乱检验针对异常组件，不是B0_R30−B0；近期地震活动的增益尚未接受专门的置乱检验。因此后续前瞻模型不纳入异常组件，但数据及负结果保留。",
    )


def _results(document: Document, figures: Path) -> None:
    _add_heading(document, "3 结果", 1)
    _add_heading(document, "3.1 30天主终点", 2)
    _add_key_box(
        document,
        "相同600,000 km²报警面积上限",
        "B0命中5/21（23.81%），B0_R30命中9/21（42.86%）；净增益为4群和19.05个百分点。21群中有5群两者都命中、4群由漏报转为命中、0群反向丢失、12群两者都漏报。",
    )
    _add_paragraph(
        document,
        "B0_R30在全部五档面积上限下均高于B0。最显著的面积效率结果出现在300,000 km²：B0_R30已命中7/21，而B0在600,000 km²仅命中5/21，相当于在预登记面积档上至少节省两档。主终点的配对Bootstrap 95%区间为+4.76至+38.10个百分点，正增益复本比例为0.9905；该结果条件于已拟合模型与21个观测震群，不构成传统显著性证明。四个不一致结果均为B0_R30新增命中、没有反向损失；作为事后敏感性检查，精确配对二项检验单侧p=0.0625、双侧p=0.125。三折命中从3/8、2/6、0/7变为5/8、4/6、0/7，三个外折均不变差，但第三折没有任何命中。",
    )
    _add_figure(
        document,
        figures / "figure_01_primary_result.png",
        "图2 30天主结果。左侧为600,000 km²一级终点；右上为五档面积—召回曲线；左下为三段时间外推；右下列出主要证据边界。",
    )

    _add_heading(document, "3.2 六模型结果与90天补充证据", 2)
    _add_table(
        document,
        ["模型", "30天命中/21", "30天召回", "相对B0", "平均实际面积（km²）"],
        [
            ("B0", "5", "23.81%", "基线", "599,447"),
            ("B0_R30", "9", "42.86%", "+4群 / +19.05 pp", "599,666"),
            ("B0_C", "6", "28.57%", "+1群 / +4.76 pp", "599,600"),
            ("B0_C_A_snapshot", "6", "28.57%", "+1群 / +4.76 pp", "599,643"),
            ("B0_C_A_dynamic", "6", "28.57%", "+1群 / +4.76 pp", "599,673"),
            ("B0_R30_C_A_dynamic", "9", "42.86%", "+4群 / +19.05 pp", "599,663"),
        ],
        widths_cm=[4.2, 2.6, 2.4, 4.0, 3.2],
    )
    _add_paragraph(
        document,
        "完整组合与B0_R30在30天主终点都命中9/21，不能把完整组合的总提升归因于异常。90天、600,000 km²次级终点中，B0_R30由4/22提高到8/22，增益为4群和18.18个百分点，95%区间为+4.55至+36.36个百分点。90天方向与主终点一致，但由于其为次级终点且起报日更少，不替代30天结论。",
    )

    _add_heading(document, "3.3 全部21群的成对变化", 2)
    _add_paragraph(
        document,
        "图3逐行展示全部21群，而不是只列出成功案例。四个新增命中分别对应2023年12月甘肃肃北M5.0、2024年3月青海杂多M5.5、2024年8月新疆库车M5.0和2024年10月新疆库车M5.5。B0原有的五个命中全部被B0_R30保留。剩余12群仍同时漏报，是模型下一步需要解释和改进的主要部分。",
    )
    _add_figure(
        document,
        figures / "figure_02_all_cluster_outcomes.png",
        "图3 全部21个规则化震群的逐一命中状态。橙色标注表示B0漏报而B0_R30命中的四个新增震群；主终点没有反向丢失。",
        width_cm=15.8,
    )

    _add_heading(document, "3.4 说明性震例", 2)
    _add_paragraph(
        document,
        "2024年10月26日新疆库车M5.5是排名变化最清楚的案例：2024年10月17日00:00（北京时间）起报时，长期背景中目标格位列第1725；加入近30天活动后升至第3。起报前30天目标75 km内有5个M4+事件，最近M4.9事件距目标约21.7 km、早于起报约4小时。该目标在起报后9.7天发生。",
    )
    _add_paragraph(
        document,
        "2023年12月1日甘肃肃北M5.0提供了不同地点的补充：2023年11月23日起报时，目标格由第3912升至第194，并进入同一面积报警区；起报前约29天，距目标14.5 km发生M5.5事件，目标在起报后约9天发生。两个震例均显示R30主要把正在活动或近期活跃的区域向前排序，而不是在没有目录信号的区域产生新的热点。",
    )
    _add_figure(
        document,
        figures / "figure_03_case_studies.png",
        "图4 两个说明性震例的局部空间排序与时间线。黄色边框为600,000 km²报警前缀在局部图中的网格；黄色圆点为起报前30天M4+事件；红星为评分后叠加的目标地震，不参与预测。",
    )
    _add_paragraph(
        document,
        "另外两个新增命中需要更谨慎地解释。2024年8月库车M5.0同样由第2146升至第3，但它与10月库车M5.5落在同一25 km网格，不能当作完全独立的空间重复。2024年3月杂多M5.5由第1669升至第153，起报前约46小时、6.7 km处已有M5.3；二者属于同一全局规则震群，因此更接近对活跃序列延续的报警，而不是对一个全新区域的发现。",
    )
    _add_figure(
        document,
        figures / "figure_04_case_rank_shifts.png",
        "图5 四个新增命中的目标格排名变化。横轴为对数名次，越靠左表示优先级越高。案例解释模型行为，总体效应仍以全部21群为准。",
    )

    _add_heading(document, "3.5 异常数据的结果", 2)
    _add_paragraph(
        document,
        "异常报告并非未使用：205期异常被重建为快照、全部异常和动态演化特征，并进入三个增量模型及时间/空间置乱。30天、600,000 km²下，B0_C、B0_C_A_snapshot和B0_C_A_dynamic均命中6/21；两个异常模型相对覆盖控制没有增加命中，动态相对快照也没有增加命中。时间和空间置乱因失败率超过预设有效性阈值而判为证据不足。因此后续前瞻模型不纳入当前异常组件，但这一结果不能外推为所有异常在所有条件下都无效。",
    )


def _discussion(document: Document) -> None:
    _add_heading(document, "4 讨论", 1)
    _add_heading(document, "4.1 近期地震活动为何有效", 2)
    _add_paragraph(
        document,
        "B0描述长期空间易发性，能覆盖反复发生地震的区域，但对短时间内突然增强的局部活动反应较弱。R30把起报前30天的M4+活动重新投影到同一空间网格，因而能抬高近期活跃区的排名。四个新增命中的近场目录均存在可见活动，库车和杂多案例尤其明显。这与ETAS及近期目录编码模型利用自激和短期聚集的基本思路一致[5–9]。",
    )
    _add_paragraph(
        document,
        "但本研究不能把所有增益解释为主震前兆。杂多案例属于同一活动序列的延续，库车两个目标位于同一网格；R30可能同时捕捉前震、余震、震群和一般短期聚集。科学上更稳妥的表述是：近期地震活动改善了目标区域的相对排序，而不是已识别出特定主震的独有前兆。",
    )

    _add_heading(document, "4.2 结果的实质意义", 2)
    _add_paragraph(
        document,
        "19.05个百分点的主终点增益不仅来自固定阈值：B0_R30在300,000 km²已超过B0在600,000 km²的命中数，说明近期活动既提高召回，也提高面积效率。对实际区域预测而言，这比在不受控面积下报告更多命中更有意义。与此同时，12/21震群仍被两个模型同时漏报，显示单靠近期地震目录不能解决大部分目标；长期构造、目录完整度、不同时间尺度和更合适的触发模型仍可能提供增量。",
    )

    _add_heading(document, "4.3 异常、复杂模型与负结果", 2)
    _add_paragraph(
        document,
        "异常表的负结果具有直接决策价值。完整组合的9/21与B0_R30完全相同，三个预登记异常增量又均为0，因此当前数据不支持继续增加异常特征或更复杂模型。相比之下，简单的B0_R30已经提供了清晰、可解释的增益。为检验异常是否超过地震聚集背景，本研究预先设置覆盖控制、时间置乱和空间置乱；这些检验未达到预设有效性阈值，因此异常没有进入真实前瞻模型。",
    )

    _add_heading(document, "4.4 主要限制", 2)
    for bullet in (
        "一级终点只有21个规则化震群，区间仍较宽；Bootstrap未包含模型重拟合和参数选择不确定性，并依赖震群可交换假设。",
        "四个新增命中全部来自前两个时间折和西北地区；第三折0/7→0/7，地域与年代泛化尚未证明。",
        "两个库车目标位于同一25 km网格，物理震群虽按时间分开，但空间重复性降低了表面上的多样性。",
        "B0_R30−B0没有完成专门的目录时间/空间置乱；现有置乱仅针对异常组件。",
        "历史原始目录采用UTC+8和发布时间等于发震时刻的冻结假定，可能使当时可用性略偏乐观。",
        "ETAS比较器因当前实现的数值资格失败而不可评价；本文只证明相对75 km KDE背景的历史增益，不能证明优于成熟ETAS。",
        "锁定测试未读取；截至2026年8月31日真实前瞻为0期，任何实际预测能力结论必须等待未来起报及其成熟真值。",
    ):
        _add_bullet(document, bullet)

    _add_heading(document, "4.5 真实前瞻验证", 2)
    _add_paragraph(
        document,
        "下一阶段不回填历史起报，也不调整模型。真实前瞻协议冻结B0、B0_R30=0.75B0+0.25R30、M4+输入、30天窗口、600,000 km²报警面积上限和规则化震群定义。首个合法起报时刻为2026年9月10日00:00（北京时间）；截至本稿冻结日尚未起报。今后每一期在预定时刻前保存不可回写的相对强度、报警网格、静态图和离线交互页，再等待30/90天真值成熟。按照CSEP式前瞻原则，只有起报时尚未知的未来地震才能提供确认性证据[4,12]。",
    )


def _conclusion(document: Document) -> None:
    _add_heading(document, "5 结论", 1)
    _add_paragraph(
        document,
        "在冻结的历史时间外推中，加入近30天地震活动的简单模型在相同600,000 km²报警面积上限下，将30天M5–6规则化震群召回从5/21提高到9/21；在300,000 km²下也已超过长期背景在600,000 km²的命中数。这是对受控报警面积下提高区域召回目标的直接历史改进。",
    )
    _add_paragraph(
        document,
        "改善主要表现为对近期活跃区域的大幅升位，符合地震聚集的统计机制。异常报告在当前冻结方案中没有带来额外主终点命中，因而没有进入前瞻模型。由于增益集中在西北、第三折无改善且样本量小，B0_R30应作为冻结候选接受未来盲检；现有结果不能证明已经实现可靠的实际地震预测。",
    )

    _add_heading(document, "数据和代码可用性", 1)
    _add_paragraph(
        document,
        "代码与冻结协议已公开于https://github.com/Justin-147/SeismoFlux，前瞻模型代码标签为v0.2.7-p1-b0-r30-code（提交c71c97790adcf33f6c8121e367317857dc8dff31）。本稿所在公开分支同时提供派生统计表、静态图、离线交互页面及其只读重建脚本。原始地震目录、异常表、断层和底图受来源与再分发限制，不随论文包附发；其规模、冻结身份和用途记录在仓库数据清单中。真实前瞻起报将以不可覆盖版本单独归档。",
    )
    _add_heading(document, "利益冲突、经费与作者贡献", 1)
    _add_paragraph(
        document,
        "利益冲突声明、经费来源和作者贡献将在投稿前由作者确认并按目标期刊格式补充。",
    )


def _references(document: Document) -> None:
    _add_heading(document, "参考文献", 1)
    references = [
        "Jordan TH. Earthquake predictability, brick by brick. Seismological Research Letters. 2006;77(1):3–6. https://doi.org/10.1785/gssrl.77.1.3",
        "Molchan GM, Kagan YY. Earthquake prediction and its optimization. Journal of Geophysical Research: Solid Earth. 1992;97(B4):4823–4838. https://doi.org/10.1029/91JB03095",
        "Zechar JD, Jordan TH. Testing alarm-based earthquake predictions. Geophysical Journal International. 2008;172(2):715–724. https://doi.org/10.1111/j.1365-246X.2007.03676.x",
        "Schorlemmer D, Gerstenberger MC, Wiemer S, Jackson DD, Rhoades DA. Earthquake likelihood model testing. Seismological Research Letters. 2007;78(1):17–29. https://doi.org/10.1785/gssrl.78.1.17",
        "Ogata Y. Statistical models for earthquake occurrences and residual analysis for point processes. Journal of the American Statistical Association. 1988;83(401):9–27. https://doi.org/10.1080/01621459.1988.10478560",
        "Ogata Y. Space-time point-process models for earthquake occurrences. Annals of the Institute of Statistical Mathematics. 1998;50:379–402. https://doi.org/10.1023/A:1003403601725",
        "Ogata Y. Statistics of earthquake activity: Models and methods for earthquake predictability studies. Annual Review of Earth and Planetary Sciences. 2017;45:497–527. https://doi.org/10.1146/annurev-earth-063016-015918",
        "Dascher-Cousineau K, Shchur O, Brodsky EE, Günnemann S. Using deep learning for flexible and scalable earthquake forecasting. Geophysical Research Letters. 2023;50:e2023GL103909. https://doi.org/10.1029/2023GL103909",
        "Zlydenko O, Elidan G, Hassidim A, et al. A neural encoder for earthquake rate forecasting. Scientific Reports. 2023;13:12350. https://doi.org/10.1038/s41598-023-38033-9",
        "Zhang Y, Zhan C, Huang Q, Sornette D. Seismically informed reference models enhance AI-based earthquake prediction systems. Journal of Geophysical Research: Solid Earth. 2024;129:e2023JB028037. https://doi.org/10.1029/2023JB028037",
        "Zhang Y, Han P, Chen H, Zhan C, Niu Y, Zhuang J, Zhu K. Incorporating non-seismicity precursors into earthquake probabilistic forecasting model. Geophysical Research Letters. 2025;52:e2025GL117972. https://doi.org/10.1029/2025GL117972",
        "Serafini F, Bayona JA, Silva F, Stockman S, Savran W, Maechling PJ, Werner MJ. A benchmark database of ten years of prospective next-day earthquake forecasts in California from the Collaboratory for the Study of Earthquake Predictability. Scientific Data. 2025;12:1501. https://doi.org/10.1038/s41597-025-05766-3",
    ]
    for index, reference in enumerate(references, start=1):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.75)
        paragraph.paragraph_format.first_line_indent = Cm(-0.75)
        paragraph.paragraph_format.space_after = Pt(3.3)
        paragraph.paragraph_format.line_spacing = 1.05
        run = paragraph.add_run(f"[{index}] {reference}")
        _set_run_font(run, size=8.8)


def build(output_path: Path, publication_root: Path) -> None:
    figures = publication_root / "figures"
    required = [
        figures / "figure_01_primary_result.png",
        figures / "figure_02_all_cluster_outcomes.png",
        figures / "figure_03_case_studies.png",
        figures / "figure_04_case_rank_shifts.png",
        figures / "figure_05_method_overview.png",
    ]
    missing = [str(path) for path in required if not path.is_file()]
    if missing:
        raise FileNotFoundError("missing publication figures: " + ", ".join(missing))

    document = Document()
    _configure_styles(document)
    _configure_sections(document)
    document.core_properties.title = "近期地震活动相对长期核密度背景的回顾性区域排序增益"
    document.core_properties.subject = "SeismoFlux retrospective scientific evaluation"
    document.core_properties.author = ""
    document.core_properties.keywords = "earthquake forecasting, recent seismicity, fixed alarm area"
    document.core_properties.comments = ""

    _title_page(document)
    _abstract(document)
    _introduction(document)
    _data_methods(document, figures)
    _results(document, figures)
    _discussion(document)
    _conclusion(document)
    _references(document)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--publication-root",
        type=Path,
        default=Path("outputs/publication/seismoflux_b0_r30_v1"),
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=Path("outputs/publication/seismoflux_b0_r30_v1/manuscript/seismoflux_manuscript_zh.docx"),
    )
    args = parser.parse_args()
    root = Path.cwd()
    publication_root = args.publication_root if args.publication_root.is_absolute() else root / args.publication_root
    output = args.output if args.output.is_absolute() else root / args.output
    build(output.resolve(), publication_root.resolve())


if __name__ == "__main__":
    main()
